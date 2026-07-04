"""
InsureIQ — Tests: Report Composer Agent (Agent 6) + report generator
Run with: pytest tests/test_report_composer.py -v

Gemini is mocked so these run without a real API key — safe for CI. The .docx
generator makes no LLM calls, so it's tested directly against real output bytes.
Key assertions: Agent 6 writes ONLY prose (numbers/flags come from upstream), the
intros are normalized so every key exists, and the report renders for sparse docs.
"""

import io

import pytest
from unittest.mock import patch, MagicMock
from docx import Document

from agents.report_composer_agent import (
    run_report_composer_agent,
    _summarize_for_prompt,
    _normalize_report_intros,
)
from tools.report_generator import generate_report


POLICY_DATA = {
    "status": "success",
    "policy_number": "POL-123",
    "insurer_name": "Acme Life",
    "policyholder_name": "Jane Doe",
    "premium_amount": "₹24,000",
    "premium_frequency": "annual",
    "policy_term_years": "20",
    "policy_start_date": "2020-01-01",
    "policy_end_date": "2040-01-01",
    "maturity_date": "2040-01-01",
    "maturity_benefit": "₹6,00,000",
    "sum_assured": "₹5,00,000",
    "nominee_name": "John Doe",
    "free_look_period_days": 15,
    "grace_period_days": 30,
}

COVERAGE_MAP = {
    "covered_events": [{"event": "Death", "conditions": "any cause after 1 year"}],
    "excluded_events": [{"event": "Suicide (year 1)", "reason": "standard exclusion"}],
    "waiting_periods": [{"condition": "Suicide clause", "duration": "12 months"}],
    "sub_limits": [{"category": "Accidental rider", "limit": "₹1,00,000"}],
    "coverage_summary": "Pays a lump sum on death; matures at 20 years.",
}

FINANCIAL_VERDICT = {
    "status": "success",
    "verdict": "NET_LOSS",
    "total_premium_paid": "₹480,000",
    "maturity_benefit": "₹600,000",
    "net_gain_loss": "₹+120,000",
    "effective_annual_return_pct": 1.12,
    "fd_benchmark_return": "₹1,539,000",
    "index_fund_benchmark_return": "₹4,632,000",
    "verdict_plain_english": "This policy grows your money slower than inflation.",
    "comparison_statement": "A fixed deposit would have returned far more.",
}

RISK_FLAGS = {
    "status": "success",
    "flags": [
        {"severity": "HIGH", "category": "Surrender penalty",
         "description": "Lose money if you exit early", "implication": "Big loss on early exit",
         "page_reference": "p5"},
        {"severity": "MEDIUM", "category": "Auto-renewal",
         "description": "Renews automatically", "implication": "You may keep paying", "page_reference": "p3"},
    ],
    "total_high": 1,
    "total_medium": 1,
    "total_low": 0,
    "overall_risk_level": "HIGH",
}

MOCK_INTROS_RESPONSE = """{
  "report_title": "InsureIQ Analysis: Life Policy — Acme Life",
  "executive_summary": "This is an endowment life policy that loses to inflation.",
  "dates_section_intro": "Here are the key dates that matter.",
  "financial_section_intro": "Now to the money question.",
  "risk_section_intro": "Watch out for these clauses.",
  "recommendation": "Consider a term plan plus an index fund instead."
}"""


def _mock_model(response_text):
    model = MagicMock()
    model.generate_content.return_value.text = response_text
    return model


class TestSummarizeForPrompt:
    def test_includes_headline_facts(self):
        summary = _summarize_for_prompt(
            POLICY_DATA, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, "life"
        )
        assert "Acme Life" in summary
        assert "NET_LOSS" in summary
        assert "HIGH" in summary
        assert "Surrender penalty" in summary

    def test_handles_empty_upstream(self):
        # Must not raise on bare dicts (no flags, no coverage).
        summary = _summarize_for_prompt({}, {}, {}, {}, "unknown")
        assert "no notable flags" in summary


class TestNormalizeReportIntros:
    def test_all_keys_present(self):
        intros = _normalize_report_intros({}, POLICY_DATA, "life")
        for key in ("report_title", "executive_summary", "dates_section_intro",
                    "financial_section_intro", "risk_section_intro", "recommendation"):
            assert key in intros

    def test_title_falls_back_when_missing(self):
        intros = _normalize_report_intros({}, POLICY_DATA, "life")
        assert "Acme Life" in intros["report_title"]
        assert "Life" in intros["report_title"]


class TestRunReportComposerAgent:
    @patch("agents.report_composer_agent.genai")
    def test_produces_valid_docx(self, mock_genai):
        mock_genai.GenerativeModel.return_value = _mock_model(MOCK_INTROS_RESPONSE)
        mock_genai.configure = MagicMock()

        result = run_report_composer_agent(
            POLICY_DATA, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, document_type="life"
        )

        assert result["status"] == "success"
        assert result["report_intros"]["report_title"].startswith("InsureIQ Analysis")
        # report_bytes must be a real, openable .docx.
        assert isinstance(result["report_bytes"], bytes)
        doc = Document(io.BytesIO(result["report_bytes"]))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "endowment life policy" in full_text  # exec summary from Gemini
        assert "Consider a term plan" in full_text    # recommendation from Gemini

    @patch("agents.report_composer_agent.genai")
    def test_invalid_json_raises(self, mock_genai):
        mock_genai.GenerativeModel.return_value = _mock_model("garbage")
        mock_genai.configure = MagicMock()

        with pytest.raises(ValueError, match="invalid JSON"):
            run_report_composer_agent(
                POLICY_DATA, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, document_type="life"
            )


class TestGenerateReport:
    def test_renders_full_report(self):
        intros = {
            "report_title": "My Report",
            "executive_summary": "Summary here.",
            "recommendation": "Cancel it.",
        }
        report = generate_report(POLICY_DATA, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, intros)
        doc = Document(io.BytesIO(report))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "My Report" in text
        assert "Cancel it." in text
        # Risk flag content shows up.
        assert "Surrender penalty" in text
        # Table content (policy number) lives in tables, not paragraphs.
        table_text = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        assert "POL-123" in table_text
        assert "Acme Life" in table_text

    def test_lapse_conditions_render_as_quick_reference(self):
        policy = dict(POLICY_DATA)
        policy["lapse_conditions"] = [
            "Premium unpaid past the 30-day grace period",
            "Material non-disclosure discovered within 2 years",
        ]
        report = generate_report(policy, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, {})
        doc = Document(io.BytesIO(report))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Conditions That Void Your Policy" in text
        assert "Premium unpaid past the 30-day grace period" in text

    def test_lapse_section_skipped_when_empty(self):
        # No lapse_conditions key → the whole block is omitted (no empty heading).
        report = generate_report(POLICY_DATA, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, {})
        doc = Document(io.BytesIO(report))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Conditions That Void Your Policy" not in text

    def test_loan_availability_row_rendered(self):
        policy = dict(POLICY_DATA)
        policy["loan_against_policy"] = "yes"
        report = generate_report(policy, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, {})
        doc = Document(io.BytesIO(report))
        table_text = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        assert "Loan against policy" in table_text
        assert "Yes" in table_text

    def test_loan_not_mentioned_row_skipped(self):
        policy = dict(POLICY_DATA)
        policy["loan_against_policy"] = "not_mentioned"
        report = generate_report(policy, COVERAGE_MAP, FINANCIAL_VERDICT, RISK_FLAGS, {})
        doc = Document(io.BytesIO(report))
        table_text = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
        assert "not_mentioned" not in table_text

    def test_renders_sparse_document_without_error(self):
        # A bare term policy: no maturity, no coverage lists, no flags.
        sparse_financial = {"verdict": "NO_MATURITY_BENEFIT", "total_premium_paid": "₹360,000",
                            "verdict_plain_english": "Pure protection, no payout."}
        report = generate_report(
            {"insurer_name": "Term Co"}, {}, sparse_financial, {}, {}
        )
        doc = Document(io.BytesIO(report))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Default title used, and the no-flags fallback line rendered.
        assert "InsureIQ" in text
        assert "No notable red flags" in text
