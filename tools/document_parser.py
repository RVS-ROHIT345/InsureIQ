"""
InsureIQ — Document Parser Utility
Handles text extraction from PDF and DOCX files.

Strategy:
1. Try pdfplumber for text-based PDFs (preserves table structure)
2. If pdfplumber returns < 100 chars, fall back to Gemini native PDF vision
   (handles scanned/image-based PDFs without any OCR library)
3. For DOCX files, use python-docx directly

This module is a pure utility — it has no agent logic.
"""

import logging
import io
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Minimum character count to consider pdfplumber extraction successful.
# Scanned PDFs often return "" or whitespace only.
MIN_EXTRACTION_LENGTH = 100


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, str]:
    """
    Extract text from a PDF file.

    Args:
        file_bytes: Raw PDF bytes (processed in memory, never written to disk)

    Returns:
        Tuple of (extracted_text, method_used)
        method_used is "pdfplumber" or "gemini_vision"

    Raises:
        ValueError: If the file cannot be parsed as a PDF
    """
    logger.info("Attempting pdfplumber extraction...")

    try:
        text = _extract_with_pdfplumber(file_bytes)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}. Will fall back to Gemini vision.")
        text = ""

    if len(text.strip()) >= MIN_EXTRACTION_LENGTH:
        logger.info(f"pdfplumber succeeded: {len(text)} characters extracted")
        return text, "pdfplumber"

    # Fallback: Gemini native PDF vision for scanned/image-based PDFs
    logger.warning(
        f"pdfplumber returned only {len(text.strip())} chars — "
        "document may be scanned. Switching to Gemini vision fallback."
    )
    text = _extract_with_gemini_vision(file_bytes)
    return text, "gemini_vision"


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, str]:
    """
    Extract text from a DOCX file.

    Args:
        file_bytes: Raw DOCX bytes (processed in memory, never written to disk)

    Returns:
        Tuple of (extracted_text, method_used)

    Raises:
        ValueError: If the file cannot be parsed as a DOCX
    """
    logger.info("Extracting text from DOCX...")

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables (premium schedules are often in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
        logger.info(f"DOCX extraction succeeded: {len(text)} characters")
        return text, "python-docx"

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {e}") from e


def _extract_with_pdfplumber(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pdfplumber.
    Better than PyPDF2 for tables — critical for premium schedule extraction.

    Args:
        file_bytes: Raw PDF bytes

    Returns:
        Extracted text as a single string
    """
    all_text = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        logger.debug(f"PDF has {len(pdf.pages)} pages")

        for page_num, page in enumerate(pdf.pages, start=1):
            # Extract regular text
            page_text = page.extract_text() or ""

            # Extract tables separately — pdfplumber handles these well
            tables = page.extract_tables()
            table_text = _format_tables(tables)

            if page_text or table_text:
                all_text.append(f"--- Page {page_num} ---")
                if page_text:
                    all_text.append(page_text)
                if table_text:
                    all_text.append(table_text)

    return "\n".join(all_text)


def _format_tables(tables: list) -> str:
    """
    Convert pdfplumber table data into readable text.

    Args:
        tables: List of tables, each a list of rows, each row a list of cells

    Returns:
        Formatted table text
    """
    if not tables:
        return ""

    formatted = []
    for table in tables:
        for row in table:
            # Filter None cells and join with pipe separator
            row_text = " | ".join(str(cell).strip() for cell in row if cell is not None)
            if row_text.strip():
                formatted.append(row_text)

    return "\n".join(formatted)


def _extract_with_gemini_vision(file_bytes: bytes) -> str:
    """
    Fallback: Send PDF bytes directly to Gemini 1.5 Pro for native PDF reading.
    Gemini can read scanned PDFs natively — no OCR library required.

    This is called only when pdfplumber extraction fails (scanned documents).

    Args:
        file_bytes: Raw PDF bytes

    Returns:
        Text extracted by Gemini

    Raises:
        RuntimeError: If Gemini API call fails
    """
    import google.generativeai as genai
    from config.settings import settings

    logger.info("Using Gemini vision fallback for scanned PDF...")

    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel(settings.GEMINI_MODEL)

    prompt = (
        "This is a scanned insurance document. Please extract ALL text from every page "
        "exactly as it appears, preserving numbers, dates, amounts, and clause numbering. "
        "Do not summarize. Return the full extracted text."
    )

    # Gemini accepts raw PDF bytes as a Part
    pdf_part = {
        "inline_data": {
            "mime_type": "application/pdf",
            "data": file_bytes,
        }
    }

    try:
        response = model.generate_content([prompt, pdf_part])
        text = response.text
        logger.info(f"Gemini vision extracted {len(text)} characters")
        return text

    except Exception as e:
        raise RuntimeError(f"Gemini vision extraction failed: {e}") from e


def detect_file_extension(filename: str) -> str:
    """
    Extract and validate file extension.

    Args:
        filename: Original filename from upload

    Returns:
        Lowercase extension without dot (e.g., "pdf", "docx")

    Raises:
        ValueError: If extension is not pdf or docx
    """
    suffix = Path(filename).suffix.lower().lstrip(".")

    if suffix not in {"pdf", "docx"}:
        raise ValueError(
            f"Unsupported file type: .{suffix}. "
            "InsureIQ accepts PDF and DOCX files only."
        )

    return suffix


def validate_file_size(file_bytes: bytes, max_bytes: int) -> None:
    """
    Enforce file size limit before processing.

    Args:
        file_bytes: File content to check
        max_bytes: Maximum allowed size in bytes

    Raises:
        ValueError: If file exceeds the size limit
    """
    size_mb = len(file_bytes) / (1024 * 1024)
    max_mb = max_bytes / (1024 * 1024)

    if len(file_bytes) > max_bytes:
        raise ValueError(
            f"File size ({size_mb:.1f} MB) exceeds the {max_mb:.0f} MB limit."
        )

    logger.debug(f"File size check passed: {size_mb:.2f} MB")
