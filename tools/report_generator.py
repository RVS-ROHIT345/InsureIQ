"""
InsureIQ — Report Generator Utility
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Called by the Report Composer Agent (Agent 6) to build the final .docx output.

This module is pure document assembly — NO LLM calls. It takes the structured
outputs of Agents 2–5 plus the plain-English section intros written by Agent 6
and lays them out into a formatted, human-readable Word report with python-docx.

Design notes:
  - Every accessor is defensive: any field may be None / missing / empty because
    upstream agents null out whatever a document didn't contain. The report must
    render cleanly for a bare-bones term policy and a 40-page ULIP alike.
  - Numbers are never computed here — they arrive already formatted from Agent 4.
"""

import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

# Severity → display colour for the risk table. Muted, print-friendly tones.
_SEVERITY_COLOR = {
    "HIGH": RGBColor(0xC0, 0x39, 0x2B),    # red
    "MEDIUM": RGBColor(0xB9, 0x77, 0x0E),  # amber
    "LOW": RGBColor(0x5D, 0x6D, 0x7E),     # grey
}

# Human labels + source keys for the policy-details table. Order matters — this is
# the sequence the reader sees.
_POLICY_FIELDS = [
    ("Policy number", "policy_number"),
    ("Insurer", "insurer_name"),
    ("Policyholder", "policyholder_name"),
    ("Nominee", "nominee_name"),
    ("Sum assured", "sum_assured"),
    ("Premium", "premium_amount"),
    ("Premium frequency", "premium_frequency"),
]

# Date/timeline fields for the "Important Dates" section.
_DATE_FIELDS = [
    ("Policy start", "policy_start_date"),
    ("Policy end", "policy_end_date"),
    ("Maturity date", "maturity_date"),
    ("Policy term (years)", "policy_term_years"),
    ("Free-look period (days)", "free_look_period_days"),
    ("Grace period (days)", "grace_period_days"),
]

_MISSING = "Not specified"


def _clean(value) -> str:
    """Render any policy value as display text, collapsing None/empty to a marker."""
    if value is None:
        return _MISSING
    text = str(value).strip()
    return text if text else _MISSING


def _format_loan_availability(value) -> str:
    """
    Render the loan_against_policy field for the Policy Details table.

    The extractor emits "yes" | "no" | "not_mentioned". We surface a clear Yes/No,
    and treat "not_mentioned"/blank as missing so the row is skipped rather than
    showing an ugly placeholder. (The actual loan *terms*, if punitive, are already
    called out in the risk section — this row is just the at-a-glance fact.)
    """
    if value is None:
        return _MISSING
    text = str(value).strip().lower()
    if text in ("yes", "y", "true", "available"):
        return "Yes"
    if text in ("no", "n", "false"):
        return "No"
    return _MISSING


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Add a two-column (label / value) table. Skips rows whose value is missing."""
    visible = [(label, val) for label, val in rows if val != _MISSING]
    if not visible:
        doc.add_paragraph("No details were found in the document for this section.")
        return

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for label, value in visible:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        # Bold the label column for scannability.
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True


def _add_bullets(doc: Document, items: list[str], empty_msg: str) -> None:
    """Add a bulleted list, or a fallback line when there's nothing to list."""
    if not items:
        doc.add_paragraph(empty_msg)
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _intro(report_intros: dict, key: str) -> str:
    """Fetch a section intro written by Agent 6; empty string if absent."""
    return (report_intros.get(key) or "").strip()


def _add_policy_section(doc: Document, policy_data: dict) -> None:
    doc.add_heading("Policy Details", level=1)
    rows = [(label, _clean(policy_data.get(key))) for label, key in _POLICY_FIELDS]
    # Loan availability needs friendly Yes/No formatting rather than raw "yes"/"not_mentioned".
    rows.append(("Loan against policy", _format_loan_availability(policy_data.get("loan_against_policy"))))
    _add_kv_table(doc, rows)


def _add_lapse_section(doc: Document, policy_data: dict) -> None:
    """
    A terse, factual "what voids your cover" quick-reference.

    Deliberately NOT an analysis — the risk section already grades the severity of
    lapse/forfeiture clauses. This block just lists the conditions in one place so a
    reader can see, at a glance, every way the policy can stop paying out. Skipped
    entirely when the extractor found no lapse conditions (no value in an empty box).
    """
    conditions = [
        str(c).strip() for c in (policy_data.get("lapse_conditions") or [])
        if str(c).strip()
    ]
    if not conditions:
        return
    doc.add_heading("Conditions That Void Your Policy", level=1)
    doc.add_paragraph(
        "Your cover can lapse or benefits be forfeited if any of the following occur. "
        "See the Red Flags section for how serious each one is."
    )
    _add_bullets(doc, conditions, "None found.")


def _add_dates_section(doc: Document, policy_data: dict, report_intros: dict) -> None:
    doc.add_heading("Important Dates", level=1)
    intro = _intro(report_intros, "dates_section_intro")
    if intro:
        doc.add_paragraph(intro)
    _add_kv_table(doc, [(label, _clean(policy_data.get(key))) for label, key in _DATE_FIELDS])


def _add_coverage_section(doc: Document, coverage_map: dict) -> None:
    doc.add_heading("What This Policy Covers", level=1)

    summary = (coverage_map.get("coverage_summary") or "").strip()
    if summary:
        doc.add_paragraph(summary)

    doc.add_heading("Covered", level=2)
    covered = [
        f"{c.get('event', '')}" + (f" — {c['conditions']}" if c.get("conditions") else "")
        for c in coverage_map.get("covered_events", []) if isinstance(c, dict)
    ]
    _add_bullets(doc, covered, "No specific covered events were extracted.")

    doc.add_heading("Excluded (NOT covered)", level=2)
    excluded = [
        f"{e.get('event', '')}" + (f" — {e['reason']}" if e.get("reason") else "")
        for e in coverage_map.get("excluded_events", []) if isinstance(e, dict)
    ]
    _add_bullets(doc, excluded, "No specific exclusions were extracted.")

    waiting = coverage_map.get("waiting_periods", [])
    if waiting:
        doc.add_heading("Waiting Periods", level=2)
        _add_bullets(
            doc,
            [f"{w.get('condition', '')}: {w.get('duration', '')}"
             for w in waiting if isinstance(w, dict)],
            "None.",
        )

    sub_limits = coverage_map.get("sub_limits", [])
    if sub_limits:
        doc.add_heading("Sub-limits", level=2)
        _add_bullets(
            doc,
            [f"{s.get('category', '')}: {s.get('limit', '')}"
             for s in sub_limits if isinstance(s, dict)],
            "None.",
        )


def _add_financial_section(doc: Document, financial_verdict: dict, report_intros: dict) -> None:
    doc.add_heading("Is It Worth the Money?", level=1)
    intro = _intro(report_intros, "financial_section_intro")
    if intro:
        doc.add_paragraph(intro)

    verdict = financial_verdict.get("verdict") or "UNKNOWN"
    verdict_para = doc.add_paragraph()
    verdict_para.add_run("Verdict: ").bold = True
    verdict_para.add_run(verdict.replace("_", " ").title())

    plain = (financial_verdict.get("verdict_plain_english") or "").strip()
    if plain:
        doc.add_paragraph(plain)

    # Metric table — only rows the calculator actually produced (protection-only
    # and insufficient-data verdicts leave most of these null).
    metric_rows = [
        ("Total premium paid", _clean(financial_verdict.get("total_premium_paid"))),
        ("Maturity benefit", _clean(financial_verdict.get("maturity_benefit"))),
        ("Net gain / loss", _clean(financial_verdict.get("net_gain_loss"))),
        ("Effective annual return (CAGR)", _fmt_pct(financial_verdict.get("effective_annual_return_pct"))),
        ("Same money in a fixed deposit", _clean(financial_verdict.get("fd_benchmark_return"))),
        ("Same money in an index fund", _clean(financial_verdict.get("index_fund_benchmark_return"))),
    ]
    _add_kv_table(doc, metric_rows)

    comparison = (financial_verdict.get("comparison_statement") or "").strip()
    if comparison:
        doc.add_paragraph(comparison)


def _fmt_pct(value) -> str:
    """Format a numeric CAGR percentage for display; missing → marker."""
    if value is None:
        return _MISSING
    try:
        return f"{float(value):g}%"
    except (TypeError, ValueError):
        return _clean(value)


def _add_risk_section(doc: Document, risk_flags: dict, report_intros: dict) -> None:
    doc.add_heading("Red Flags & Fine Print", level=1)
    intro = _intro(report_intros, "risk_section_intro")
    if intro:
        doc.add_paragraph(intro)

    overall = risk_flags.get("overall_risk_level") or "NONE"
    summary_para = doc.add_paragraph()
    summary_para.add_run("Overall risk level: ").bold = True
    overall_run = summary_para.add_run(overall)
    overall_run.bold = True
    if overall in _SEVERITY_COLOR:
        overall_run.font.color.rgb = _SEVERITY_COLOR[overall]
    summary_para.add_run(
        f"  ({risk_flags.get('total_high', 0)} high, "
        f"{risk_flags.get('total_medium', 0)} medium, "
        f"{risk_flags.get('total_low', 0)} low)"
    )

    flags = risk_flags.get("flags", [])
    if not flags:
        doc.add_paragraph("No notable red flags were found in this document.")
        return

    for flag in flags:
        if not isinstance(flag, dict):
            continue
        severity = str(flag.get("severity", "LOW")).upper()
        heading = doc.add_paragraph()
        sev_run = heading.add_run(f"[{severity}] ")
        sev_run.bold = True
        if severity in _SEVERITY_COLOR:
            sev_run.font.color.rgb = _SEVERITY_COLOR[severity]
        heading.add_run(flag.get("category") or "Uncategorized").bold = True

        if flag.get("description"):
            doc.add_paragraph(str(flag["description"]))
        if flag.get("implication"):
            impl = doc.add_paragraph()
            impl.add_run("What it means for you: ").italic = True
            impl.add_run(str(flag["implication"]))
        if flag.get("page_reference"):
            ref = doc.add_paragraph()
            ref_run = ref.add_run(f"(Location: {flag['page_reference']})")
            ref_run.font.size = Pt(9)
            ref_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_recommendation_section(doc: Document, report_intros: dict) -> None:
    recommendation = _intro(report_intros, "recommendation")
    if not recommendation:
        return
    doc.add_heading("Our Recommendation", level=1)
    doc.add_paragraph(recommendation)


def generate_report(
    policy_data: dict,
    coverage_map: dict,
    financial_verdict: dict,
    risk_flags: dict,
    report_intros: dict,
) -> bytes:
    """
    Assemble all agent outputs into a formatted .docx report.

    Args:
        policy_data: Output from Policy Extractor Agent (Agent 2)
        coverage_map: Output from Coverage Analyzer Agent (Agent 3)
        financial_verdict: Output from Financial Evaluator Agent (Agent 4)
        risk_flags: Output from Risk Flag Agent (Agent 5)
        report_intros: Section introductions from Report Composer Agent (Agent 6)

    Returns:
        DOCX file as bytes (caller is responsible for returning to user)
    """
    policy_data = policy_data or {}
    coverage_map = coverage_map or {}
    financial_verdict = financial_verdict or {}
    risk_flags = risk_flags or {}
    report_intros = report_intros or {}

    logger.info("Generating .docx report")

    doc = Document()

    # ── Title + executive summary ─────────────────────────────────────────────
    title = doc.add_heading(
        report_intros.get("report_title") or "InsureIQ Policy Analysis", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    exec_summary = (report_intros.get("executive_summary") or "").strip()
    if exec_summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(exec_summary)

    # ── Body sections (order chosen for a reader deciding whether to keep it) ──
    _add_policy_section(doc, policy_data)
    _add_dates_section(doc, policy_data, report_intros)
    _add_lapse_section(doc, policy_data)
    _add_coverage_section(doc, coverage_map)
    _add_financial_section(doc, financial_verdict, report_intros)
    _add_risk_section(doc, risk_flags, report_intros)
    _add_recommendation_section(doc, report_intros)

    # ── Footer disclaimer ─────────────────────────────────────────────────────
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    dis_run = disclaimer.add_run(
        "Generated by InsureIQ. This is an automated analysis to aid understanding, "
        "not professional financial or legal advice. Verify all figures against your "
        "original policy document."
    )
    dis_run.italic = True
    dis_run.font.size = Pt(8)
    dis_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buffer = io.BytesIO()
    doc.save(buffer)
    report_bytes = buffer.getvalue()
    logger.info(f"Report generated — {len(report_bytes)} bytes")
    return report_bytes
