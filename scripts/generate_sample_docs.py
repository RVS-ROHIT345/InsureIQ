"""
InsureIQ — Sample Document Generator (bootstrap / gap-filler)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Purpose: guarantee that a checkout is never left WITHOUT a runnable fixture for
the health / car / home document types. It emits simple, self-describing specimen
policies (PDF via ReportLab, DOCX via python-docx) plus one deliberately-sparse
"cover note" edge case, using the exact filenames the test suite expects.

IMPORTANT — this is a fallback, not the source of truth. The fixtures actually
committed under sample_docs/ are richer, hand-crafted documents (e.g. the SENTINEL
health-cum-savings endowment that carries a genuine maturity benefit, and the
NORTHSTAR limited-pay life plans). Those hand-crafted docs are what exercise the
nuanced downstream behaviour — financial routing on a real maturity benefit,
unusual-vs-standard risk calibration, etc. The specimens this script produces are
intentionally plainer and will NOT reproduce those cases (in particular, the
generated health doc is a 1-year mediclaim with no maturity benefit). It also does
not generate the `life` fixtures at all.

Because of that, this script is strictly NON-DESTRUCTIVE: it only writes a file
that is genuinely missing and never overwrites an existing one. That is what keeps
it from silently replacing a hand-crafted fixture with an inferior stand-in. If you
truly want to regenerate a specimen, delete the file first, then run this.

Every document is stamped as a specimen and contains enough type-specific
vocabulary (see config.settings.DOC_TYPE_KEYWORDS) that the keyword classifier
resolves the type without an LLM call — which is what makes the fixtures testable
offline.

Usage:
    python scripts/generate_sample_docs.py    # fills in only genuinely-missing files
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SAMPLE_DIR = PROJECT_ROOT / "sample_docs"

SPECIMEN_NOTE = (
    "Specimen document generated for InsureIQ software testing purposes. "
    "Not a real insurance policy. All names, numbers, and clauses are fictitious."
)

# ── Document definitions ──────────────────────────────────────────────────────
# Each doc is: title, subtitle, a declarations table (list of [label, value] rows),
# and a list of sections. A section is (heading, [paragraphs], table-or-None).
# Type-specific vocabulary is woven in naturally so the keyword classifier locks
# onto the right type from the first page.

HEALTH_DOC = {
    "stem": "sample_health_insurance_policy",
    "title": "CarePlus Health Insurance Co.",
    "subtitle": "Comprehensive Mediclaim Hospitalization Plan — Policy Document",
    "declarations": [
        ["Policy Number", "CPH/2026/HLT/0098231"],
        ["Product Name", "CarePlus Family Floater Health Insurance"],
        ["Policyholder", "Specimen Insured"],
        ["Sum Insured", "Rs. 10,00,000 (Ten Lakhs)"],
        ["Annual Premium", "Rs. 24,500 (including GST)"],
        ["Policy Term", "1 year (annually renewable)"],
        ["Date of Commencement", "01 April 2026"],
        ["Renewal Date", "01 April 2027"],
    ],
    "sections": [
        (
            "1. Scope of Cover",
            [
                "This health insurance policy is a mediclaim contract covering "
                "hospitalization and medical expenses incurred by the insured for "
                "illness or injury during the policy period, subject to the terms "
                "and conditions set out below.",
                "Cashless treatment is available at any network hospital. Where "
                "treatment is taken at a non-network hospital, expenses are "
                "reimbursed on submission of original bills.",
            ],
            None,
        ),
        (
            "2. Benefits Covered",
            [
                "The following medical expenses are payable up to the sum insured:",
            ],
            [
                ["Benefit", "Limit"],
                ["In-patient hospitalization", "Up to sum insured"],
                ["Room rent (normal)", "1% of sum insured per day"],
                ["ICU room rent", "2% of sum insured per day"],
                ["Pre-hospitalization", "60 days"],
                ["Post-hospitalization", "90 days"],
                ["Day-care surgical benefit", "Covered — 540 listed procedures"],
                ["Critical illness lump sum", "Rs. 5,00,000 (optional rider)"],
                ["Annual health check-up", "Once per policy year"],
            ],
        ),
        (
            "3. Waiting Periods",
            [
                "Any pre-existing disease declared at inception is covered only "
                "after a continuous waiting period of 36 months.",
                "A general initial waiting period of 30 days applies from the "
                "date of commencement, except for accidental injury.",
                "Specified illnesses (cataract, hernia, joint replacement) carry "
                "a waiting period of 24 months.",
            ],
            None,
        ),
        (
            "4. Exclusions",
            [
                "The policy does not cover: cosmetic or aesthetic treatment; "
                "dental treatment unless arising from accident; expenses for "
                "any condition arising from self-inflicted injury or attempted "
                "suicide; treatment for alcohol or drug abuse; and expenses that "
                "are not medically necessary.",
                "Room rent charges exceeding the eligible limit are subject to "
                "proportionate deduction on the entire hospital bill.",
            ],
            None,
        ),
        (
            "5. Claims & Renewal",
            [
                "For cashless treatment, intimate the network hospital's insurance "
                "desk at least 48 hours before a planned admission, or within 24 "
                "hours of an emergency admission.",
                "The policy is renewable for life provided renewal premium is paid "
                "before the renewal date. A grace period of 30 days is allowed, "
                "during which cover is not available for any claim.",
                "A no-claim bonus of 10% of the sum insured (max 50%) is added for "
                "every claim-free year.",
            ],
            None,
        ),
    ],
}

CAR_DOC = {
    "stem": "sample_car_insurance",
    "title": "DriveSafe General Insurance Ltd.",
    "subtitle": "Comprehensive Private Car Package Motor Insurance Policy",
    "declarations": [
        ["Policy Number", "DSG/2026/MOT/5540127"],
        ["Product Name", "DriveSafe Comprehensive Motor Insurance"],
        ["Insured", "Specimen Insured"],
        ["Vehicle", "2024 Model, Private Car, Petrol"],
        ["Insured Declared Value (IDV)", "Rs. 8,50,000"],
        ["Own Damage Premium", "Rs. 18,200"],
        ["Third Party Premium", "Rs. 7,890"],
        ["Total Premium", "Rs. 30,745 (including GST)"],
        ["Policy Period", "01 April 2026 to 31 March 2027"],
    ],
    "sections": [
        (
            "1. Nature of Cover",
            [
                "This motor insurance policy provides comprehensive cover for the "
                "insured vehicle, combining own damage cover with the statutory "
                "third party liability cover required under the Motor Vehicles Act.",
                "The insured declared value (IDV) shown above is the agreed value "
                "of the vehicle and is the maximum amount payable in the event of "
                "total loss or theft.",
            ],
            None,
        ),
        (
            "2. Section I — Own Damage",
            [
                "The company will indemnify the insured against loss of or damage "
                "to the vehicle arising from:",
            ],
            [
                ["Peril", "Covered"],
                ["Accident / collision damage", "Yes"],
                ["Fire, explosion, self-ignition", "Yes"],
                ["Theft / burglary", "Yes"],
                ["Flood, storm, natural calamity", "Yes"],
                ["Riot, strike, malicious act", "Yes"],
                ["In-transit by road, rail, air", "Yes"],
            ],
        ),
        (
            "3. Section II — Third Party Liability",
            [
                "The company will indemnify the insured against legal liability for "
                "third party bodily injury or death, and for third party property "
                "damage up to Rs. 7,50,000, arising out of the use of the vehicle.",
            ],
            None,
        ),
        (
            "4. No Claim Bonus",
            [
                "A no claim bonus (NCB) is allowed on renewal for each claim-free "
                "year, ranging from 20% after the first year to 50% after five "
                "consecutive claim-free years. The NCB is forfeited entirely if any "
                "claim is made during the policy period.",
            ],
            None,
        ),
        (
            "5. Add-on Covers & Exclusions",
            [
                "Optional add-ons in force: zero depreciation cover, engine "
                "protection, and 24x7 roadside assistance.",
                "The policy excludes: normal wear and tear and depreciation; "
                "mechanical or electrical breakdown; damage while driving under the "
                "influence of alcohol or drugs; damage while driving without a valid "
                "licence; and consequential loss.",
                "A compulsory deductible of Rs. 1,000 applies to every own damage "
                "claim.",
            ],
            None,
        ),
    ],
}

HOME_DOC = {
    "stem": "sample_home_insurance",
    "title": "HearthGuard Property Insurance Co.",
    "subtitle": "Homeowner Dwelling & Contents Home Insurance Policy",
    "declarations": [
        ["Policy Number", "HGP/2026/HOM/7781540"],
        ["Product Name", "HearthGuard Comprehensive Home Insurance"],
        ["Homeowner", "Specimen Insured"],
        ["Structure (Dwelling) Coverage", "Rs. 75,00,000"],
        ["Contents Coverage", "Rs. 15,00,000"],
        ["Annual Premium", "Rs. 12,300 (including GST)"],
        ["Policy Term", "1 year"],
        ["Policy Period", "01 April 2026 to 31 March 2027"],
    ],
    "sections": [
        (
            "1. About This Cover",
            [
                "This home insurance policy is a property insurance contract that "
                "protects the homeowner's building structure and its contents "
                "against the insured perils listed below, for the policy term shown "
                "in the schedule.",
                "Structure coverage insures the physical building; contents coverage "
                "insures household goods, furniture, and personal effects.",
            ],
            None,
        ),
        (
            "2. Insured Perils",
            [
                "The dwelling coverage and contents coverage respond to loss or "
                "damage caused by:",
            ],
            [
                ["Peril", "Structure", "Contents"],
                ["Fire damage & lightning", "Yes", "Yes"],
                ["Flood damage & inundation", "Yes", "Yes"],
                ["Storm, cyclone, tempest", "Yes", "Yes"],
                ["Earthquake (optional add-on)", "Yes", "Yes"],
                ["Burglary & housebreaking", "No", "Yes"],
                ["Escape of water / burst pipes", "Yes", "Yes"],
            ],
        ),
        (
            "3. Additional Benefits",
            [
                "Alternative accommodation / loss of rent is payable up to Rs. "
                "50,000 per month for a maximum of 6 months where the home is "
                "rendered uninhabitable by an insured peril.",
                "Public liability cover of Rs. 5,00,000 is included for third party "
                "injury or property damage occurring at the insured premises.",
            ],
            None,
        ),
        (
            "4. Exclusions",
            [
                "The policy does not cover: loss or damage due to wear and tear, "
                "gradual deterioration, or lack of maintenance; damage to property "
                "left unoccupied for more than 30 consecutive days; loss of cash, "
                "jewellery, or valuables unless specifically declared; and damage "
                "caused by war, invasion, or nuclear perils.",
            ],
            None,
        ),
        (
            "5. Claims & Conditions",
            [
                "In the event of burglary or malicious damage, the insured must "
                "report the incident to the police and to the company within 24 "
                "hours. Claims are settled on a reinstatement basis for the "
                "structure and on an indemnity (depreciated) basis for contents.",
                "Under-insurance is subject to the condition of average: if the sum "
                "insured is less than the actual value at risk, claims are reduced "
                "proportionately.",
            ],
            None,
        ),
    ],
}

# A deliberately sparse edge case: a one-page motor cover note with only the
# barest facts. It still classifies as "car" (two+ keywords) but omits most of
# the fields downstream agents look for — exercising their defensive handling.
EDGE_COVER_NOTE = {
    "stem": "sample_car_insurance_cover_note",
    "title": "DriveSafe General Insurance Ltd.",
    "subtitle": "Motor Insurance Cover Note (Provisional)",
    "declarations": [
        ["Cover Note No.", "DSG/CN/2026/00417"],
        ["Vehicle", "Private Car"],
        ["Insured Declared Value (IDV)", "Rs. 6,00,000"],
        ["Premium Received", "Rs. 22,100"],
        ["Valid From", "01 April 2026"],
    ],
    "sections": [
        (
            "Provisional Cover",
            [
                "This cover note confirms that comprehensive motor insurance "
                "(own damage plus third party liability) is in force for the "
                "vehicle described above, pending issue of the full policy "
                "document. Terms, conditions, and exclusions of the company's "
                "standard private car package policy apply.",
            ],
            None,
        ),
    ],
}

DOCS = [HEALTH_DOC, CAR_DOC, HOME_DOC, EDGE_COVER_NOTE]


# ── DOCX renderer ─────────────────────────────────────────────────────────────
def render_docx(doc_def: dict, out_path: Path) -> None:
    doc = DocxDocument()

    title = doc.add_heading(doc_def["title"], level=0)
    doc.add_heading(doc_def["subtitle"], level=1)
    note = doc.add_paragraph()
    run = note.add_run(SPECIMEN_NOTE)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_heading("Policy Declarations", level=2)
    decl = doc.add_table(rows=0, cols=2)
    decl.style = "Light Grid Accent 1"
    for label, value in doc_def["declarations"]:
        cells = decl.add_row().cells
        cells[0].text = label
        cells[1].text = value

    for heading, paragraphs, table in doc_def["sections"]:
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            doc.add_paragraph(para)
        if table:
            header, *rows = table
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(header):
                t.rows[0].cells[i].text = h
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val

    doc.save(str(out_path))


# ── PDF renderer ──────────────────────────────────────────────────────────────
def render_pdf(doc_def: dict, out_path: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"], fontSize=18, alignment=TA_CENTER
    )
    subtitle_style = ParagraphStyle(
        "DocSubtitle", parent=styles["Heading2"], alignment=TA_CENTER, fontSize=12
    )
    note_style = ParagraphStyle(
        "Note", parent=styles["Normal"], fontSize=7,
        textColor=colors.grey, alignment=TA_CENTER,
    )
    heading_style = styles["Heading2"]
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6
    )

    flow = [
        Paragraph(doc_def["title"], title_style),
        Spacer(1, 4),
        Paragraph(doc_def["subtitle"], subtitle_style),
        Spacer(1, 4),
        Paragraph(SPECIMEN_NOTE, note_style),
        Spacer(1, 10),
        Paragraph("Policy Declarations", heading_style),
    ]

    decl_table = Table(doc_def["declarations"], colWidths=[65 * mm, 100 * mm])
    decl_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef3fb")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    flow.append(decl_table)
    flow.append(Spacer(1, 10))

    for heading, paragraphs, table in doc_def["sections"]:
        flow.append(Paragraph(heading, heading_style))
        for para in paragraphs:
            flow.append(Paragraph(para, body_style))
        if table:
            n_cols = len(table[0])
            width = 165 * mm
            t = Table(table, colWidths=[width / n_cols] * n_cols)
            t.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe5f5")),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            flow.append(t)
        flow.append(Spacer(1, 8))

    SimpleDocTemplate(
        str(out_path), pagesize=A4,
        topMargin=18 * mm, bottomMargin=18 * mm,
        leftMargin=20 * mm, rightMargin=20 * mm,
        title=doc_def["subtitle"], author="InsureIQ sample generator",
    ).build(flow)


def main() -> None:
    SAMPLE_DIR.mkdir(exist_ok=True)

    written, skipped = [], []
    for doc_def in DOCS:
        stem = doc_def["stem"]
        # This generator only emits the DOCX cover note; the committed corpus may
        # additionally carry a hand-crafted PDF cover note, which we never touch.
        formats = ["docx"] if stem.endswith("cover_note") else ["docx", "pdf"]
        for fmt in formats:
            out_path = SAMPLE_DIR / f"{stem}.{fmt}"
            # Strictly non-destructive: an existing fixture is authoritative (it may be
            # a richer hand-crafted document) and is never overwritten by a specimen.
            if out_path.exists():
                skipped.append(out_path.name)
                continue
            if fmt == "docx":
                render_docx(doc_def, out_path)
            else:
                render_pdf(doc_def, out_path)
            written.append(f"{out_path.name} ({out_path.stat().st_size} bytes)")

    print("InsureIQ sample document generator")
    print("-" * 40)
    for name in written:
        print(f"  wrote   {name}")
    for name in skipped:
        print(f"  kept    {name} (already present — never overwritten)")
    if not written:
        print("  Nothing to write. All fixtures already present.")


if __name__ == "__main__":
    main()
